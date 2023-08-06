import json
import copy
import xlwt
import datetime

from docx import Document
from django.shortcuts import render
from django.db import connections, connection
from decimal import Decimal
from django.http import HttpResponse
from django.conf import settings
from operator import indexOf
from datetime import date, datetime
from django.db.models import Count
from django.utils.encoding import force_str
from django.utils.decorators import available_attrs
from django.shortcuts import resolve_url
from django.utils.six.moves.urllib.parse import urlparse
from django.core.serializers.json import DjangoJSONEncoder
from docxtpl import DocxTemplate
import collections
import unicodedata

try:
    from collections import OrderedDict
except ImportError:
    from ordereddict import OrderedDict

try:
    from functools import wraps
except ImportError:
    from django.utils.functional import wraps  # Python 2.4 fallback.

from models import Chart, ChartAuthGroup, ChartAuthPermission, ChartGroup,\
    ChartMetric, ScheduledReport, ScheduledReportParam, ReportRecipients,\
    Group


def request_passes_test(test_func, login_url=None, redirect_field_name=None):
    """
    Decorator for views that checks that the request passes the given test,
    redirecting to the log-in page if necessary. The test should be a callable
    that takes the request object and returns True if the request passes.
    """

    def decorator(view_func):
        @wraps(view_func, assigned=available_attrs(view_func))
        def _wrapped_view(request, *args, **kwargs):
            if test_func(request):
                return view_func(request, *args, **kwargs)
            elif hasattr(settings, 'HIDASH_SETTINGS') and 'api_redirector' in settings.HIDASH_SETTINGS:
                return settings.HIDASH_SETTINGS['api_redirector'](request)
            else:
                return HttpResponse("User Not Authorized", status=401)
        return _wrapped_view
    return decorator


def authenticate_url(request, url_params=""):
    if hasattr(settings, 'HIDASH_SETTINGS') and 'api_authenticator' in settings.HIDASH_SETTINGS:
        return settings.HIDASH_SETTINGS['api_authenticator'](request, url_params)
    else:
        return True


def _get_group_filters(request):
    group_id = request.GET.get('group', 'default')
    filters = []
    filter_value_list = [{}]
    if hasattr(settings, 'HIDASH_SETTINGS') and 'filter_values' in settings.HIDASH_SETTINGS:
        filter_value_list = []
        filter_value_list.append(settings.HIDASH_SETTINGS['filter_values'](request))
    group_params = Group.objects.filter(name=group_id)[0].chartparameter_set.all().annotate(null_priority=Count('priority')).order_by('-null_priority', 'priority')
    for param in group_params:
        if param.parameter_type == 0:
            filter_values = [filter.get(str(param.parameter_name), [{'value': 1, 'name': 'No values provided'}]) for filter in filter_value_list]
            type = 'dropdown'
            filters.append({
                            'name': str(param.parameter_name),
                            'type': type,
                            'filter_values': filter_values[0],
                            'grid_width': param.grid_width
                            })
        elif param.parameter_type == 2 or param.parameter_type == 3:
            parameter_name = str(param.parameter_name).split(',')
            start_date = [filter.get(parameter_name[0], {'value': None}) for filter in filter_value_list]
            end_date = [filter.get(parameter_name[1], {'value': None}) for filter in filter_value_list]
            type = 'daterange'
            filter_values = [start_date[0]['value'], end_date[0]['value']]
            if param.parameter_type == 3:
                type = 'periodicDatePicker'
                if hasattr(settings, 'HIDASH_SETTINGS') and 'periodic_config' in settings.HIDASH_SETTINGS:
                    filter_values = settings.HIDASH_SETTINGS['periodic_config']
            filters.append({
                            'name': parameter_name,
                            'type': type,
                            'grid_width': param.grid_width,
                            'filter_value': filter_values
                            })

        else:
            filter_value = [filter.get(str(param.parameter_name), {'value': None}) for filter in filter_value_list]
            type = 'datetime'
            filters.append({
                            'name': str(param.parameter_name),
                            'type': type,
                            'grid_width': param.grid_width,
                            'filter_value': filter_value[0]['value']
                            })
    return filters


def _group_reports_as_json(request):
    group_id = request.GET.get('group', 'default')
    params = _augment_params(request)
    charts = _load_charts()
    data = []
    for chart_id, chart in charts.iteritems():
        if group_id in chart.group:
            if check_permissions(chart, request) and check_groups(chart, request):
                chartdata = {}
                handler = _handler_selector(chart)
                chartdata['chart_data'] = handler(chart,
                                                  chart.query,
                                                  params)
                chartdata['chart_id'] = chart_id
                if chart.lib is not None:
                    chartdata['handler_type'] = 'googlechart'
                    if chart.lib == 0:
                        chartdata['handler_type'] = 'highcharts'
                else:
                    chartdata['handler_type'] = 'hidash'
                chartdata['chart_type'] = chart.chart_type
                chartdata['description'] = chart.description
                chartdata['height'] = chart.height
                chartdata['grid_width'] = chart.grid_width
                data.append(chartdata)
    return data


@request_passes_test(lambda u: authenticate_url(u, "reports_as_json_api"), login_url=None, redirect_field_name=None)
def dispatch_group_reports_as_json(request):
    data = _group_reports_as_json(request)
    return HttpResponse(content=json.dumps(data, cls=DjangoJSONEncoder),
                        content_type="application/json")


@request_passes_test(lambda u: authenticate_url(u, "reports_api"), login_url=None, redirect_field_name=None)
def dispatch_group_reports(request):
    return render(request, 'reports.html')


@request_passes_test(lambda u: authenticate_url(u, "groups_api"), login_url=None, redirect_field_name=None)
def dispatch_groups(request):
    """
    Renders the hidash groups template and provides it a context with all the groups and the
    charts associated with them
    """
    groups = []
    all_groups = ChartGroup.objects.all().prefetch_related('chart')
    for group in all_groups:
        if not any(hidash_group.get('group_name', None) == group.group.name for hidash_group in groups):
            groups.append({'group_name': str(group.group.name), 'charts': []})

    for group in all_groups:
        group_index = next(index for (index, d) in enumerate(groups) if d["group_name"] == str(group.group.name))
        groups[group_index]['charts'].append(str(group.chart))

    return render(request, 'hidash-groups.html',
                  {'hidash_groups': groups})


@request_passes_test(lambda u: authenticate_url(u, "chart_configurations_api"), login_url=None, redirect_field_name=None)
def dispatch_chart_configurations(request):
    return HttpResponse(content=json.dumps(_get_chart_congiguration(request), cls=DjangoJSONEncoder),
                        content_type="application/json")


@request_passes_test(lambda u: authenticate_url(u, "group_filters_api"), login_url=None, redirect_field_name=None)
def dispatch_group_filters_as_json(request):
    return HttpResponse(content=json.dumps(_get_group_filters(request), cls=DjangoJSONEncoder),
                 content_type="application/json")


def _get_chart_congiguration(request):
    charts = _load_charts()
    chart_configs = []
    for chart_id, chart in charts.iteritems():
        if check_permissions(chart, request) and check_groups(chart, request):
            chart_config = {
                            'chart': chart_id,
                            'description': chart.description,
                            'height': chart.height,
                            'grid_width': chart.grid_width,
                            'chart_type': chart.chart_type
                            }
            if chart.lib is not None:
                chart_config['handler_type'] = 'googlechart'
                if chart.lib == 0:
                    chart_config['handler_type'] = 'highcharts'
            else:
                chart_config['handler_type'] = 'hidash'
            chart_configs.append(chart_config)
    return chart_configs


@request_passes_test(lambda u: authenticate_url(u, "reports_as_excel_api"), login_url=None, redirect_field_name=None)
def dispatch_xls(request, chart_id):
    '''
    Function to render reports in spreadsheet format available for download
    '''
    chart_id = chart_id.split('.')[0]
    params = _augment_params(request)
    wb = xlwt.Workbook()
    ws = wb.add_sheet(chart_id)
    font_style = xlwt.easyxf('font: name Times New Roman, color-index green, bold on;align: wrap on', num_format_str='#,##0.00')
    charts = _load_charts()
    for key, chart in charts.iteritems():
        if key == chart_id:
            if check_permissions(chart, request) and check_groups(chart, request):
                cols = []
                cols.append(chart.dimension.asdict())
                cols.extend(map(lambda c: c.asdict(), chart.metrics))
                for col in cols:
                    ws.col(indexOf(cols, col)).width = int(13 * 260)

                with connections[chart.database].cursor() as cursor:
                    cursor.execute(chart.query, params)
                    for desc in cursor.description:
                        ws.write(0, indexOf(cursor.description, desc), desc[0],
                                 font_style)

                    for db_row in cursor:
                        for col_index, chart_col in enumerate(cols):
                            value = db_row[col_index]
                            value = _convert_to_type(value, chart_col['type'])
                            ws.write(indexOf(cursor, db_row) + 1, col_index,
                                     value)
                response = HttpResponse(content_type='application/vnd.ms-excel')
                response['Content-Disposition'] = 'attachment; filename=Report.xls'
                wb.save(response)
                return response
            else:
                return HttpResponse("User Not Authorized", status=401)


def check_groups(chart, request):
    user_groups = []
    for group in request.user.groups.all():
        user_groups.append(group.name)
    for group_name in chart.groups_list:
        if group_name not in user_groups:
            return False
    return True


def check_permissions(chart, request):
    for permission in chart.permissions_list:
        if not request.user.has_perm(permission):
            return False
    return True


def _handler_selector(chart):
    # TODO: Find a way to get rid of if else
    if chart.chart_type == 'widget':
        handler = globals()['widget']
    elif chart.lib == 1:
        handler = globals()['multiple_series_row']
        if chart.chart_type == "MapChart":
            handler = globals()['google_map_chart']
        elif chart.separator is None:
            handler = globals()['report']
            if chart.dimension.id == "extract":
                handler = globals()['col_to_series_handler']
            elif len(chart.metrics) >= 1:
                handler = globals()['default_handler']
    elif chart.lib == 0:
        handler = globals()['multiple_series_row_highcharts_handler']
        if chart.separator is None:
            if chart.chart_type == 'table':
                handler = globals()['tabular_data_handler']
            elif chart.dimension.id == "extract":
                handler = globals()['col_to_series_highcharts_handler']
            elif len(chart.metrics) == 1:
                handler = globals()['single_series_highcharts_handler']
            elif len(chart.metrics) >= 1:
                handler = globals()['multiple_series_column_highcharts_handler']
    return handler


@request_passes_test(lambda u: authenticate_url(u, "charts_api"), login_url=None,
                     redirect_field_name=None)
def dispatch_chart(request, chart_id):
    """
    This view renders the chart data in desirable format to the controller
    """
    chart_id = chart_id.split('.')[0]
    params = _augment_params(request)
    charts = _load_charts()
    for key, chart in charts.iteritems():
        if key == chart_id:
            if check_permissions(chart, request) and check_groups(chart,
                                                                  request):
                handler = _handler_selector(chart)
                data = handler(chart, chart.query, params)
                return HttpResponse(content=json.dumps(data, cls=DjangoJSONEncoder),
                                    content_type="application/json")
            else:
                return HttpResponse("User Not Authorized", status=401)


@request_passes_test(lambda u: authenticate_url(u, "default_api"), login_url=None,
                     redirect_field_name=None)
def index(request):
    return render(request, 'index.html')


@request_passes_test(lambda u: authenticate_url(u, "default_api"), login_url=None,
                     redirect_field_name=None)
def hidash_index(request):
    return render(request,'hidash_index.html')


def report(chart, query, params=None):
    data = {}
    data['rows'] = rows = []
    data['cols'] = cols = []
    data['chart_type'] = 'Table'
    cursor = connections[chart.database].cursor()
    cursor.execute(chart.query, params)
    for desc in cursor.description:
        cols.append({'type': 'string', 'label': desc[0]})
    for db_row in cursor:
        row_list = []
        for col_index, chart_col in enumerate(cols):
            value = db_row[col_index]
            if isinstance(value, basestring):
                # If value contains a non english alphabet
                temp = {"v": value.encode('utf-8')}
            else:
                temp = {"v": str(value)}
            row_list.append(temp)
        rows.append({"c": row_list})
    return data


def google_map_chart(chart, query, params=None):
    data = {}
    data['rows'] = rows = []
    data['cols'] = cols = []
    cursor = connections[chart.database].cursor()
    cursor.execute(chart.query, params)
    for desc in cursor.description:
        if indexOf(cursor.description, desc) < 2:
            field_type = 'number'
        else:
            field_type = 'string'
        cols.append({'type': field_type, 'label': desc[0]})
    for db_row in cursor:
        row_list = []
        for col_index, chart_col in enumerate(cols):
            value = db_row[col_index]
            temp = {"v": str(value)}
            row_list.append(temp)
        rows.append({"c": row_list})
    return data


def widget(chart, query, params=None):
    """
    Handles single widget data
    """
    cursor = connections[chart.database].cursor()
    cursor.execute(chart.query, params)
    data = {}
    data['type'] = "widget"
    data['widget_data'] = list(cursor.fetchall())
    return data


def multiple_series_row(chart, query, params=None):

    """
    Handles the multiple series data when the series name are
    to be extracted from the rows of query set
    """
    data = {}
    data['cols'] = cols = []
    data['rows'] = rows = []
    cursor = connections[chart.database].cursor()
    cursor.execute(chart.query, params)
    columns = []
    dimension_values = []
    separator_values = []
    temp_val = []

    for column_description in cursor.description:
        columns.append(column_description[0])

    index_of_dimension = columns.index(chart.dimension.id)
    index_of_separator = columns.index(chart.separator)
    index_of_metric = columns.index(chart.metrics[0].id)

    for db_row in cursor:
        if db_row[index_of_dimension] not in dimension_values:
            dimension_values.append(db_row[index_of_dimension])
            rows.append({"c": [{"v":  _convert_to_type(db_row[index_of_dimension],
                                                       chart.dimension.type)}]})
        if db_row[index_of_separator] not in separator_values:
            separator_values.append(db_row[index_of_separator])
            temp_val.append({"v": _fill_missing_values(chart.metrics[0].type)})

    for row in rows:
        row['c'].extend(copy.deepcopy(temp_val))

    for db_row in cursor:
        for row in rows:
            if row['c'][0]['v'] == _convert_to_type(db_row[index_of_dimension],
                                                    chart.dimension.type):
                index = 1 + separator_values.index(db_row[index_of_separator])
                rows[indexOf(rows, row)]['c'][index]['v'] = _convert_to_type(db_row[index_of_metric], chart.metrics[0].type)

    cols.append(chart.dimension.asdict())
    for series in separator_values:
        cols.append({"id": series, "type": chart.metrics[0].type,
                     "label": series})

    return data


def default_handler(chart, query, params=None):
    data = {}
    data['rows'] = rows = []
    data['cols'] = cols = []
    cols.append(chart.dimension.asdict())
    cols.extend(map(lambda c: c.asdict(), chart.metrics))
    data['chart_type'] = chart.chart_type
    cursor = connections[chart.database].cursor()
    cursor.execute(chart.query, params)
    for db_row in cursor:
        row_list = []
        for col_index, chart_col in enumerate(cols):
            row_list.append({"v": _convert_to_type(db_row[col_index],
                                                   chart_col['type'])})
        rows.append({"c": row_list})
    return data


def single_series_highcharts_handler(chart, query, params=None):
    data = {'data': []}
    cols = []
    cols.append({'type': 'string', 'id': chart.dimension,
                 'label': chart.dimension})
    cols.extend(map(lambda c: c.asdict(), chart.metrics))
    cursor = connections[chart.database].cursor()
    cursor.execute(chart.query, params)
    for db_row in cursor:
        row_list = []
        for col_index, chart_col in enumerate(cols):
            row_list.append(_convert_to_type(db_row[col_index],
                                             chart_col['type']))
        data['data'].append(row_list)
        data['name'] = cols[1]['label']
    return [data]


def col_to_series_highcharts_handler(chart, query, params=None):
    data = []
    cursor = connections[chart.database].cursor()
    cursor.execute(chart.query, params)
    for db_row in cursor:
        temp = {'data': [], 'name': str(db_row[0])}
        for col_name in cursor.description:
            if indexOf(cursor.description, col_name) != 0:
                temp_list = []
                temp_list.append(col_name[0])
                temp_list.append(float(db_row[indexOf(cursor.description,
                                                      col_name)]))
                temp['data'].append(temp_list)
        data.append(temp)
    if len(data) == 0:
        data = [{"data": []}]
    return data


def col_to_series_handler(chart, query, params=None):
    data = {}
    data['cols'] = cols = []
    data['chart_type'] = chart.chart_type
    cursor = connections[chart.database].cursor()
    cursor.execute(chart.query, params)
    rows = [None] * (len(cursor.description) - 1)
    for i in range(0, len(rows)):
        temp = {"c": []}
        rows[i] = temp
    for index, desc in enumerate(cursor.description):
        if index == 0:
            cols.append({'type': 'string', 'label': desc[0]})
        else:
            rows[index-1]["c"].append({"v": desc[0]})

    for db_row in cursor:
        for index, val in enumerate(db_row):
            if index == 0:
                cols.append({'type': 'number', 'label': val})
            else:
                rows[index-1]["c"].append({"v": float(val)})

    data['rows'] = rows
    return data


def multiple_series_column_highcharts_handler(chart, query, params=None):
    """
    Handles the multiple series chart when the table has series names in
    the columns
    """
    data = []
    cols = []
    cols.append(chart.dimension)
    cols.extend(map(lambda c: c.asdict(), chart.metrics))
    cursor = connections[chart.database].cursor()
    cursor.execute(chart.query, params)
    columns = []
    for column_description in cursor.description:
        columns.append(column_description[0])
    for i in range(len(cursor.description)-1):
        data.append({'data': []})
    for db_row in cursor:
        for col_index, chart_col in enumerate(cols):
            data_list = []
            if col_index is not 0:
                data_list.append(db_row[0])
                data_list.append(_convert_to_type(db_row[columns.index(str(chart_col['label']))],
                                                  chart_col['type']))
                data[col_index-1]['data'].append(copy.deepcopy(data_list))
                data[col_index-1]['name'] = chart_col['label']
    return data


def multiple_series_row_highcharts_handler(chart, query, params=None):
    """
    Handles the multiple series data when the series name are
    to be extracted from the rows of query set
    """
    chart_data = []
    cursor = connections[chart.database].cursor()
    cursor.execute(chart.query, params)
    dimension_values = []
    separator_values = []
    columns = []
    for column_description in cursor.description:
        columns.append(column_description[0])
    for db_row in cursor:
        if [db_row[columns.index(chart.dimension.id)], _fill_missing_values(chart.metrics[0].type)] not in dimension_values:
            dimension_values.append([db_row[columns.index(chart.dimension.id)],
                                     _fill_missing_values(chart.metrics[0].type)])
        if db_row[columns.index(chart.separator)] not in separator_values:
            separator_values.append(db_row[columns.index(chart.separator)])

    for series in separator_values:
        single_chart_data = {'data': []}
        single_chart_data['data'] = copy.deepcopy(dimension_values)
        single_chart_data['name'] = series
        chart_data.append(single_chart_data)
    cursor = connections[chart.database].cursor()
    cursor.execute(chart.query, params)

    for dbrow in cursor:
        for single_series_obj in chart_data:
            if dbrow[columns.index(chart.separator)] == single_series_obj['name']:
                for data in single_series_obj['data']:
                    if [dbrow[columns.index(chart.dimension.id)], _fill_missing_values(chart.metrics[0].type)] == data:
                        data[1] = dbrow[columns.index(chart.metrics[0].id)]
    return chart_data


def tabular_data_handler(chart, query, params=None):
    """
    Handler used when tabular data is required on the front end
    """
    table_data = {'headers': [], 'rows': []}
    cursor = connections[chart.database].cursor()
    cursor.execute(chart.query, params)
    for header in cursor.description:
        table_data['headers'].append(header[0])
    for db_row in cursor:
        new_db_row = []
        for row_data in db_row:
            new_db_row.append(str(row_data))
        table_data['rows'].append(list(new_db_row))
    return table_data


def _fill_missing_values(data_type):
    if data_type == "string":
        return ""
    elif data_type == "number":
        return 0


def _convert_to_type(value, type_desired):
    if not value:
        return_value = None
    elif type_desired == 'string':
        if isinstance(value, basestring):
            return_value = value
        elif isinstance(value, (date, datetime)):
            return_value = "%s.0" % str(value)
        else:
            return_value = str(value)
    elif type_desired == 'number':
        return_value = _coerce_number(value)
    elif type_desired == 'timeofday':
        return_value = _to_time_of_day(value)
    elif type_desired == 'date':
        return_value = 'Date(' + str(value.year) + ', ' + str(value.month) + ', ' + str(value.day) + ')'
    else:
        assert False, "Invalid column type %s" % type_desired
    return return_value


def _coerce_number(possibly_number, default=0):
    if isinstance(possibly_number, Decimal):
        return float(possibly_number)
    if isinstance(possibly_number, (int, float)):
        return possibly_number
    try:
        if isinstance(possibly_number, basestring):
            return float(possibly_number.strip())
        else:
            return int(possibly_number)
    except:
        return default


def _to_time_of_day(str_or_int):
    if isinstance(str_or_int, int):
        seconds = str_or_int
        minutes = seconds / 60
        seconds = seconds % 60
        hours = minutes / 60
        minutes = minutes % 60
        return (hours, minutes, seconds, 0)
    elif isinstance(str_or_int, float):
        seconds = str_or_int
        minutes = seconds / 60
        seconds = seconds % 60
        hours = minutes / 60
        minutes = minutes % 60
        return (hours, minutes, seconds, 0)
    elif isinstance(str_or_int, basestring):
        tokens = str_or_int.split(":")
        if len(tokens) > 1:
            hour = _coerce_number(tokens[0], 0)
            minute = _coerce_number(tokens[1], 0)
            return (hour, minute, 0, 0)
        else:
            return (0, 0, 0, 0)
    else:
        assert False


class Dimension(object):
    '''Represents a <dimension> from charts model'''

    dimension_types = ['string', 'number', 'timeofday', 'date']

    def __init__(self, dimension_id, dim_type):
        assert dim_type in Dimension.dimension_types, 'Unsupported dimension type %s' % dim_type
        self.id = dimension_id.replace(" ", "_") if dimension_id else ''
        self.type = dim_type

    def asdict(self):
        return {"id": self.id, "label": self.id, "type": self.type}


class Metric(object):
    '''Represents a <dimension> from charts model'''

    metric_types = ['string', 'number', 'timeofday', 'date']

    def __init__(self, id, metric_type="number"):
        assert metric_type in Metric.metric_types, 'Unsupported metric type %s' % metric_type
        self.id = id
        self.type = metric_type

    def asdict(self):
        return {"id": self.id, "label": self.id, "type": self.type}


class ChartData(object):
    '''
    Represents a chart from the hidash models
    '''
    def __init__(self, name, database, chart_group_list, lib, permissions_list,
                 groups_list, query, dimension, metrics, separator,
                 chart_type, description, grid_width, height):

        self.id = name
        self.database = database
        self.query = query
        self.chart_type = chart_type
        self.group = chart_group_list
        self.permissions_list = permissions_list
        self.groups_list = groups_list
        self.lib = lib
        self.dimension = dimension
        self.metrics = metrics
        self.description = description
        self.grid_width = grid_width
        self.height = height
        if separator:
            self.separator = separator
        else:
            self.separator = None


def _parse_charts():
    parsed_charts = OrderedDict()
    charts = Chart.objects.all().annotate(null_priority=Count('priority')).order_by('-null_priority', 'priority')
    chart_metrics = ChartMetric.objects.all()
    chart_auth_groups = ChartAuthGroup.objects.all()
    chart_auth_permissions = ChartAuthPermission.objects.all()
    chart_groups = ChartGroup.objects.all()
    for chart in list(charts):
        current_metrics = []
        permissions_list = []
        auth_groups_list = []
        chart_group_list = []
        separator = chart.separator
        for chart_metric in chart_metrics:
            if chart_metric.chart.id == chart.id:
                current_metrics.append(Metric(chart_metric.metric))
        for groups in chart_groups:
            if groups.chart.id == chart.id:
                chart_group_list.append(groups.group.name)
        for chart_auth_permission in chart_auth_permissions:
            if chart_auth_permission.chart.id == chart.id:
                permissions_list.append(chart_auth_permission.permission)
        for chart_auth_group in chart_auth_groups:
            if chart_auth_group.chart.id == chart.id:
                auth_groups_list.append(chart_auth_group.auth_group)
        # TODO: Later on add support for multiple databases and dimension and metric type
        parsed_charts[chart.name] = ChartData(chart.name, 'default', chart_group_list,
                                              chart.library,
                                              permissions_list, auth_groups_list,
                                              chart.query, Dimension(chart.dimension, 'string'),
                                              current_metrics, separator,
                                              chart.chart_type, chart.description,
                                              chart.grid_width, chart.height)
    return parsed_charts


def _augment_params(request):
    processors = []
    if hasattr(settings, 'HIDASH_SETTINGS') and 'parameter_processors' in settings.HIDASH_SETTINGS:
        processors = settings.HIDASH_SETTINGS['parameter_processors']
    params = request.GET.copy()
    for processor in processors:
        params.update(processor(request))
    return params


def _load_charts():
    return _parse_charts()


@request_passes_test(lambda u: authenticate_url(u, "groups_api"), login_url=None, redirect_field_name=None)
def download_as_docx(request):
    """
    Populates the placeholders in the template returns a docx content type 
    """
    group = request.GET.get('group', None)
    params = _augment_params(request)
    data = _get_template_data(group, params)
    template_path = Group.objects.filter(name=group)[0].template
    if template_path == '':
        return HttpResponse('No template found. Please upload a template for this group to use this functionality')
    actual_template_path = template_path
    if hasattr(settings, 'HIDASH_SETTINGS') and 'template_context_provider' in settings.HIDASH_SETTINGS:
        context_provider = settings.HIDASH_SETTINGS['template_context_provider']
        data = context_provider(request, data)
    if hasattr(settings, 'HIDASH_SETTINGS') and 'template_downloader' in settings.HIDASH_SETTINGS:
        template_downloader = settings.HIDASH_SETTINGS['template_downloader']
        template_info = template_downloader(request, template_path)
        if template_info.get('template_path'):
            actual_template_path = template_info.get('template_path', None)
    template_data = convert(data)
    doc = DocxTemplate(str(actual_template_path))
    doc.render(template_data)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=report.docx'
    doc.save(response)
    return response


def _get_template_data(group, params=None):
    """
    Returns the data for all the charts in a group
    """
    get_templated_reports()
    reports_data = {}
    reports_data['data'] = []
    report_grouped_data =  ChartGroup.objects.all().filter(group__name=group).order_by('chart__priority')
    for report in report_grouped_data:
        report_dict = {}
        report_dict['chart_id'] = report.chart.name
        report_dict['chart_data'] = {}
        report_dict['chart_data']['rows'] = []
        cursor = connection.cursor()
        cursor.execute(report.chart.query, params)
        rows = cursor.fetchall()
        for row in rows:
            report_dict['chart_data']['rows'].append(list(row))
        header_meta = cursor.description
        report_header = []
        for header in header_meta:
            report_header.append(header[0])
        report_dict['chart_data']['columns'] = report_header
        reports_data[report.chart.name.replace(' ','_').lower()] = report_dict
        reports_data['data'].append(report_dict)
    return reports_data


def get_email_reports(reports):
    '''Returns the formatted reports which needs to be mailed on a timely basis'''
    reports_data = {}
    reports_data['data'] = []
    reports_to_be_emailed = reports
    scheduled_report_param = ScheduledReportParam.objects.filter(scheduled_report = reports_to_be_emailed.id)
    report_grouped_data =  ChartGroup.objects.all().filter(group__name=reports_to_be_emailed.group).order_by('chart__priority')
    for report in report_grouped_data:
        report_dict = {}
        report_dict['chart_id'] = report.chart.name
        report_dict['chart_data'] = {}
        report_dict['chart_data']['rows'] = []
        cursor = connection.cursor()
        bind_param = _sql_query_param_interpreter(scheduled_report_param)
        cursor.execute(report.chart.query, bind_param)
        rows = cursor.fetchall()
        for row in rows:
            report_dict['chart_data']['rows'].append(list(row))
        header_meta = cursor.description
        report_header = []
        for header in header_meta:
            report_header.append(header[0])
        report_dict['chart_data']['columns'] = report_header
        reports_data['data'].append(report_dict)
    for k in bind_param:
        reports_data[k]=bind_param[k]
    return reports_data


def get_email_data_and_recipients():
    cur_dt = datetime.now()
    scheduled_reports = ScheduledReport.objects.prefetch_related('group').filter(next_run_at__lt = cur_dt, template='')
    all_report_recipients = ReportRecipients.objects.all()
    response = {}
    for report in scheduled_reports:
        response[report.id] = get_email_reports(report)
        recipients = all_report_recipients.filter(report = report.id)
        email_to = ''
        response[report.id]['email_subject'] = report.email_subject
	response[report.id]['email_message'] = report.email_message
        for recipient in recipients:
            email_to += recipient.email_address+','
        response[report.id]['email_to'] = email_to
        report.last_run_at = datetime.now()
        report.save()
    return response


def get_templated_reports():
    cur_dt = datetime.now()
    scheduled_reports = ScheduledReport.objects.prefetch_related('group').exclude(template__exact='').filter(next_run_at__lt = cur_dt)
    response = {}
    all_report_recipients = ReportRecipients.objects.all()
    response = {}
    for report in scheduled_reports:
        response[report.id] = get_formatted_reports(report)
        recipients = all_report_recipients.filter(report = report.id)
        email_to = ''
        response[report.id]['file'] = report.template.name
        response[report.id]['email_subject'] = report.email_subject
	response[report.id]['email_message'] = report.email_message
        for recipient in recipients:
            email_to += recipient.email_address+','
        response[report.id]['email_to'] = email_to
        report.last_run_at = datetime.now()
        report.save()
    return response

def get_formatted_reports(reports):
    '''Returns the formatted reports which needs to be mailed on a timely basis'''
    reports_data = {}
    reports_data['data'] = []
    reports_to_be_emailed = reports
    scheduled_report_param = ScheduledReportParam.objects.filter(scheduled_report = reports_to_be_emailed.id)
    report_grouped_data =  ChartGroup.objects.filter(group__name=reports_to_be_emailed.group).order_by('chart__priority')
    for report in report_grouped_data:
        report_dict = {}
        report_dict['chart_id'] = report.chart.name
        report_dict['chart_data'] = {}
        report_dict['chart_data']['rows'] = []
        cursor = connection.cursor()
        bind_param = _sql_query_param_interpreter(scheduled_report_param)
        cursor.execute(report.chart.query, bind_param)
        rows = cursor.fetchall()
        for row in rows:
            report_dict['chart_data']['rows'].append(list(row))
        header_meta = cursor.description
        report_header = []
        for header in header_meta:
            report_header.append(header[0])
        report_dict['chart_data']['columns'] = report_header
        reports_data[report.chart.name.replace(' ','_').lower()] = report_dict
        reports_data['data'].append(report_dict)
    for k in bind_param:
        reports_data[k]=bind_param[k]
    return reports_data


def _sql_query_param_interpreter(report_param):
    bind_params = {}
    for param in report_param:
        if param.is_parameter_value_sql_function:
            query = '''SELECT {}'''.format(param.parameter_value)
            cursor = connection.cursor()
            cursor.execute(query)
            rows = cursor.fetchall()
            bind_params[param.parameter_name] = str(rows[0][0])
        else:
            bind_params[param.parameter_name] = param.parameter_value
    return bind_params


def create_doc_file(data, header_image_path, doc_name):
    document = Document()
    document.add_picture(header_image_path)
    for t in data['data']:
       document.add_heading(t['chart_id'].title(), level=1)
       table = document.add_table(rows=1, cols=len(t['chart_data']['columns']),style='Light Grid Accent 1')
       heading_cells = table.rows[0].cells
       for index in range(len(t['chart_data']['columns'])):
           heading_cells[index].text = str(t['chart_data']['columns'][index])
       for d in t['chart_data']['rows']:
           row_cells = table.add_row().cells
           for col in range(len(t['chart_data']['columns'])):
                if isinstance(d[col], basestring):
                   # If value contains a non english alphabet
                    row_cells[col].text = unicodedata.normalize('NFKD', d[col]).encode('ascii', 'ignore')
                else:
                    row_cells[col].text = str(d[col])
    document.save(doc_name)


def create_templated_doc(data, template_path, saved_file_path):
    data = convert(data)
    doc = DocxTemplate(str(template_path))
    doc.render(data)
    doc.save(saved_file_path)


def convert(data):
   if isinstance(data, basestring):
       if isinstance(data, unicode):
	   data = unicodedata.normalize('NFKD', data).encode('ascii', 'ignore').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"','&quot;').replace("'", '&#39;')      
	   return data
       elif isinstance(data, str):
	   return data.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;').replace("'", '&#39;')
       else:
           return str(data).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;').replace("'", '&#39;')
   elif isinstance(data, collections.Mapping):
       return dict(map(convert, data.iteritems()))
   elif isinstance(data, collections.Iterable):
       return type(data)(map(convert, data))
   else:
       return data




